import os
import logging
import asyncio
import json
import io
import re
import hashlib
from fastapi import FastAPI, WebSocket, WebSocketDisconnect, UploadFile, File, Form, Request,HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from dotenv import load_dotenv
from pydantic import BaseModel
from typing import List, Optional
from datetime import datetime, timedelta

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- SUPABASE & STRIPE ---
from supabase import create_client, Client
import stripe

# --- DEEPGRAM ---
from deepgram import (
    DeepgramClient,
    DeepgramClientOptions,
    LiveTranscriptionEvents,
    LiveOptions,
)

# --- DOCUMENT PARSERS ---
from pypdf import PdfReader  # Reverted back to your original pypdf
from docx import Document

# --- INTERNAL MODULES ---
from core.brain import Brain
from core.llm import stream_completion

# --- NEW: AI CLIENT FOR COACH ---
from groq import Groq 

load_dotenv()

print("👀 PYTHON SEES THESE VARIABLES:", list(os.environ.keys()))

# Setup Logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s | %(levelname)s | %(message)s')
logger = logging.getLogger("backend")

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["X-ATS-Score", "X-Missing-Keywords", "X-Items-Removed"]
)

# --- IN-MEMORY SESSION MANAGEMENT ---
user_sessions = {}

def get_brain_for_user(user_id: str):
    """Retrieves or creates a unique Brain instance for a specific user."""
    if user_id not in user_sessions:
        logger.info(f"🧠 Creating new Brain instance for user: {user_id}")
        user_sessions[user_id] = Brain()
    return user_sessions[user_id]


# --- API KEYS & CLIENTS SETUP ---
raw_api_key = os.getenv("DEEPGRAM_API_KEY")
DEEPGRAM_API_KEY = raw_api_key.strip() if raw_api_key else None

if not DEEPGRAM_API_KEY:
    logger.error("❌ Deepgram API Key missing! Check .env file.")

SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_KEY = os.getenv("SUPABASE_SERVICE_KEY")

if SUPABASE_URL and SUPABASE_KEY:
    supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)
    logger.info("🟢 Supabase Admin Client Connected")
else:
    logger.error("❌ Supabase URL or Service Key missing from .env!")

# STRIPE SETUP
stripe.api_key = os.getenv("STRIPE_SECRET_KEY")

# Initialize the AI Client for the Coach feature
# Ensure GROQ_API_KEY is in your .env
coach_llm_client = Groq(api_key=os.getenv("GROQ_API_KEY"))


# --- HELPER FUNCTIONS ---
def extract_text_from_file(file_content: bytes, filename: str) -> str:
    text = ""
    try:
        if filename.endswith(".pdf"):
            reader = PdfReader(io.BytesIO(file_content))
            for page in reader.pages:
                text += page.extract_text() + "\n"
        elif filename.endswith(".docx"):
            doc = Document(io.BytesIO(file_content))
            for para in doc.paragraphs:
                text += para.text + "\n"
        else:
            text = file_content.decode("utf-8")
    except Exception as e:
        logger.error(f"Error parsing file: {e}")
        return "Error reading resume."
    return text

def get_difficulty_instruction(level: str):
    if level == "Easy":
        return "Ask standard behavioral questions (e.g., 'Tell me about yourself'). Be encouraging."
    elif level == "Medium":
        return "Ask standard technical questions relevant to the resume. Be professional."
    elif level == "Hard":
        return "Ask very difficult, deep technical questions and edge cases. Be skeptical."
    return "Ask standard questions."

def add_markdown_paragraph(doc_element, text: str, style=None):
    """
    Helper function to parse **bold** markdown and add it to a python-docx paragraph.
    """
    p = doc_element.add_paragraph(style=style)
    # Split the text by the markdown bold syntax **
    parts = re.split(r'(\*\*.*?\*\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # It's bold! Remove the asterisks and add as bold run
            clean_text = part[2:-2]
            p.add_run(clean_text).bold = True
        else:
            # Normal text
            p.add_run(part)
    return p

def create_optimized_word_doc(ai_data: dict, original_text: str) -> io.BytesIO:
    doc = Document()
    
    for section in doc.sections:
        section.top_margin, section.bottom_margin = Inches(1), Inches(1)
        section.left_margin, section.right_margin = Inches(1), Inches(1)

    # 1. Contact Info
    name = ai_data.get("contact_info", {}).get("name", "Name Not Found")
    contact_str = ai_data.get("contact_info", {}).get("contact_string", "Contact Info Not Found")
    
    header_p = doc.add_paragraph()
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = header_p.add_run(name)
    name_run.bold = True
    name_run.font.size = Pt(24)
    name_run.font.color.rgb = RGBColor(31, 73, 125)
    
    contact = doc.add_paragraph(contact_str)
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line.add_run("_" * 50).font.color.rgb = RGBColor(180, 180, 180)

    # 2. Summary
    if ai_data.get("summary"):
        doc.add_heading('PROFESSIONAL PROFILE', level=1).runs[0].font.color.rgb = RGBColor(31, 73, 125)
        add_markdown_paragraph(doc, ai_data["summary"])

    # 3. Skills
    if ai_data.get("skills"):
        doc.add_heading('CORE COMPETENCIES & SKILLS', level=1).runs[0].font.color.rgb = RGBColor(31, 73, 125)
        for skill in ai_data["skills"]:
            add_markdown_paragraph(doc, skill, style='List Bullet')

    # 4. Experience (The AI will no longer drop jobs here!)
    if ai_data.get("experience"):
        doc.add_heading('PROFESSIONAL EXPERIENCE', level=1).runs[0].font.color.rgb = RGBColor(31, 73, 125)
        for job in ai_data["experience"]:
            p = doc.add_paragraph()
            if job.get("title"): p.add_run(f"{job['title']}").bold = True
            if job.get("company"): p.add_run(f" | {job['company']}")
            if job.get("dates"): p.add_run(f" | {job['dates']}")
            
            for bullet in job.get("bullets", []):
                add_markdown_paragraph(doc, bullet, style='List Bullet')

    # 5. Education & Certifications
    if ai_data.get("certifications"):
        doc.add_heading('CERTIFICATIONS & REQUIREMENTS', level=1).runs[0].font.color.rgb = RGBColor(31, 73, 125)
        for cert in ai_data["certifications"]:
            add_markdown_paragraph(doc, cert, style='List Bullet')

    # 6. Tier 3 Gap Bridger
    if ai_data.get("gap_bridger_project") and ai_data.get("gap_bridger_bullets"):
        doc.add_heading('TECHNICAL PROJECTS & UPSKILLING', level=1).runs[0].font.color.rgb = RGBColor(31, 73, 125)
        p = doc.add_paragraph()
        p.add_run(f"{ai_data['gap_bridger_project']}").bold = True
        for bullet in ai_data["gap_bridger_bullets"]:
            add_markdown_paragraph(doc, bullet, style='List Bullet')

    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return doc_io

# --- DATA MODELS ---
class CheckoutRequest(BaseModel):
    token: str
    return_url: str

class CoachReply(BaseModel):
    history: List[dict]
    resume_text: str  
    job_description: str
    user_answer: str
    difficulty: str

class SyncTimeReq(BaseModel):
    user_id: str
    minutes_to_deduct: int


# ==========================================
#         REST ENDPOINTS (CORE & BILLING)
# ==========================================

@app.post("/sync-time")
async def sync_time(req: SyncTimeReq):
    """Officially deducts minutes from the database during Coach sessions."""
    try:
        curr_res = await asyncio.to_thread(
            lambda: supabase.table("user_credits").select("balance_minutes").eq("user_id", req.user_id).single().execute()
        )
        curr_bal = curr_res.data.get("balance_minutes", 0)
        
        if curr_bal > 0:
            new_bal = curr_bal - req.minutes_to_deduct
            await asyncio.to_thread(
                lambda: supabase.table("user_credits").update({"balance_minutes": new_bal}).eq("user_id", req.user_id).execute()
            )
            return {"status": "success", "new_balance": new_bal}
        return {"status": "insufficient_funds"}
    except Exception as e:
        logger.error(f"Sync time error: {e}")
        return {"error": str(e)}

@app.post("/submit-context")
async def submit_context(
    user_id: str = Form(...),  
    resume: UploadFile = File(...), 
    job_description: str = Form(default="")
):
    """Used for the Live Copilot Context Setup"""
    try:
        user_brain = get_brain_for_user(user_id)
        content = await resume.read()
        resume_text = extract_text_from_file(content, resume.filename)
        
        user_brain.set_context(resume_text, job_description)
        logger.info(f"✅ Context updated for User {user_id}. Resume length: {len(resume_text)}")
        return {"status": "success"}
    except Exception as e:
        logger.error(f"Upload failed: {e}")
        return {"status": "error", "message": str(e)}

@app.post("/create-checkout-session")
async def create_checkout_session(req: CheckoutRequest):
    try:
        user_res = await asyncio.to_thread(supabase.auth.get_user, req.token)
        user_id = user_res.user.id
        user_email = user_res.user.email

        session = stripe.checkout.Session.create(
            payment_method_types=['card'],
            line_items=[{
                'price': 'price_1T4pnAEFLEZAHwelrTvPR1Os', 
                'quantity': 1,
            }],
            mode='payment',
            success_url=f"{req.return_url}/?success=true",
            cancel_url=f"{req.return_url}/?canceled=true",
            customer_email=user_email,
            client_reference_id=user_id, 
        )
        return {"url": session.url}
    except Exception as e:
        logger.error(f"Stripe error: {str(e)}")
        return {"error": str(e)}

@app.post("/webhook")
async def stripe_webhook(request: Request):
    payload = await request.body()
    sig_header = request.headers.get("stripe-signature")
    webhook_secret = os.getenv("STRIPE_WEBHOOK_SECRET")

    try:
        event = stripe.Webhook.construct_event(payload, sig_header, webhook_secret)
    except ValueError:
        return {"error": "Invalid payload"}, 400
    except stripe.error.SignatureVerificationError:
        return {"error": "Invalid signature"}, 400

    if event['type'] == 'checkout.session.completed':
        session = event['data']['object']
        user_id = session.get('client_reference_id') 

        if user_id:
            logger.info(f"💰 Payment success for {user_id}. Adding 60 mins.")
            try:
                curr_res = await asyncio.to_thread(
                    lambda: supabase.table("user_credits").select("balance_minutes").eq("user_id", user_id).single().execute()
                )
                curr_bal = curr_res.data.get("balance_minutes", 0)
                new_bal = curr_bal + 60
                await asyncio.to_thread(
                    lambda: supabase.table("user_credits").update({"balance_minutes": new_bal}).eq("user_id", user_id).execute()
                )
                logger.info(f"✅ Balance updated to {new_bal}")
            except Exception as e:
                logger.error(f"❌ Failed to update Supabase: {e}")

    return {"status": "success"}
# ==========================================
#         AI Optimizer Endpoint
# ==========================================
free_tier_usage = {}

# ==========================================
# ENDPOINT: OPTIMIZE RESUME
# ==========================================
@app.post("/optimize")
async def optimize_resume(
    request: Request,
    job_description: str = Form(...),
    tier: int = Form(...), 
    resume_text: str = Form(""),
    resume_file: UploadFile = File(None),
    user_id: str = Form("guest") 
):
    # 1. Billing & Access Control Logic
    curr_bal = 0
    now = datetime.now()

    if tier == 1:
        # --- FREE TIER LOGIC (Tier 1) ---
        identifier = request.client.host if user_id == "guest" else user_id
        
        # Check if they have used their 1 free optimization in the last 24 hours
        if identifier in free_tier_usage:
            last_used = free_tier_usage[identifier]
            if now < last_used + timedelta(days=1):
                # They are locked out!
                if user_id == "guest":
                    raise HTTPException(status_code=429, detail="Guest limit reached (1 per day). Please sign up for another free optimization!")
                else:
                    raise HTTPException(status_code=402, detail="Free limit reached (1 per day). Please top up your minutes to continue optimizing.")
        
        # If they pass, record their usage for today
        free_tier_usage[identifier] = now
        print(f"Free Tier 1 used by: {identifier}")

        # Permanently log guest metrics in Supabase
        if user_id == "guest":
            try:
                hashed_ip = hashlib.sha256(identifier.encode()).hexdigest()
                await asyncio.to_thread(
                    lambda: supabase.table("guest_usage_logs").insert({
                        "hashed_ip": hashed_ip,
                        "feature_used": "resume_optimizer_tier_1"
                    }).execute()
                )
            except Exception as e:
                logger.error(f"Failed to log guest usage: {e}")
                # We do NOT raise an error here. Let them have their resume even if logging fails.

    else:
        # --- PAID TIER LOGIC (Tiers 2 & 3) ---
        if user_id == "guest":
            raise HTTPException(status_code=401, detail="Please log in to use advanced tiers.")
            
        cost_map = {2: 25, 3: 50}
        minutes_to_deduct = cost_map.get(tier, 25)

        try:
            curr_res = await asyncio.to_thread(
                lambda: supabase.table("user_credits").select("balance_minutes").eq("user_id", user_id).single().execute()
            )
            curr_bal = curr_res.data.get("balance_minutes", 0)
            
            if curr_bal < minutes_to_deduct:
                raise HTTPException(status_code=402, detail=f"You need {minutes_to_deduct} minutes for this tier. Please top up.")
                
            new_bal = curr_bal - minutes_to_deduct
            await asyncio.to_thread(
                lambda: supabase.table("user_credits").update({"balance_minutes": new_bal}).eq("user_id", user_id).execute()
            )
        except Exception as e:
            logger.error(f"Failed to check/deduct balance: {e}")
            raise HTTPException(status_code=500, detail="Database error.")

    # 2. Extract Text
    final_resume_text = resume_text
    if resume_file:
        content = await resume_file.read()
        final_resume_text = extract_text_from_file(content, resume_file.filename)

    # ==========================================
    # STEP 1: THE EXTRACTOR (Data Fidelity Only)
    # ==========================================
    extractor_prompt = f"""
    You are an expert data extraction algorithm. 
    Your ONLY job is to convert the following raw resume text into a perfectly structured JSON object.
    
    CRITICAL RULES:
    1. Do NOT rewrite, summarize, or optimize anything. Copy the text exactly as it appears.
    2. You MUST identify EVERY distinct job in the "PROFESSIONAL EXPERIENCE" section. Look for patterns like "Title | Company | Location | Dates". 
    3. If there are multiple jobs, you MUST create an object for each one in the `experience` array. Do NOT combine them.
    
    RAW RESUME TEXT:
    {final_resume_text[:4000]}
    
    Return ONLY a valid JSON object matching this schema:
    {{
      "contact_info": {{"name": "...", "contact_string": "..."}},
      "summary": "...",
      "skills": ["...", "..."],
      "experience": [
          {{
            "title": "...",
            "company": "...",
            "dates": "...",
            "bullets": ["...", "..."]
          }}
      ],
      "certifications": ["...", "..."]
    }}
    """

    try:
        extractor_response = coach_llm_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "system", "content": extractor_prompt}],
            temperature=0.0, 
            response_format={"type": "json_object"}
        )
        extracted_data = json.loads(extractor_response.choices[0].message.content)
        print(f"\n[STEP 1] Jobs Extracted: {len(extracted_data.get('experience', []))}\n")
        
    except Exception as e:
        logger.error(f"Extractor failed: {e}")
        # SECURE REFUND: Only refund if it's a paid tier and NOT a guest!
        if user_id != "guest" and tier > 1:
            await asyncio.to_thread(lambda: supabase.table("user_credits").update({"balance_minutes": curr_bal}).eq("user_id", user_id).execute())
        raise HTTPException(status_code=500, detail="Failed to parse resume.")

    # ==========================================
    # STEP 2: THE OPTIMIZER (Rewriting Only)
    # ==========================================
    tier_instructions = ""
    if tier == 1:
        tier_instructions = "- TIER 1: Make minor tweaks to inject keywords from the Job Description. Fix grammar. Do NOT alter the core meaning of the bullets."
    elif tier == 2:
        tier_instructions = "- TIER 2: Heavily rewrite the bullet points under Experience. Use strong action verbs, infer high-level responsibilities, and quantify achievements."
    elif tier == 3:
        tier_instructions = "- TIER 3: Do everything in Tier 2. PLUS, identify 1-2 critical skills missing from their resume that the JD requires. Generate a realistic 'Independent Project'."

    optimizer_prompt = f"""
    You are an Elite Career Coach. 
    I have already extracted the candidate's resume into a structured JSON format. 
    Your job is to optimize this JSON structure based on the Job Description provided.

    Job Description:
    {job_description[:3000]}

    Extracted Resume JSON:
    {json.dumps(extracted_data)}

    CRITICAL RULES:
    1. You MUST process and return EVERY single job present in the `experience` array I provided you. If I gave you 2 jobs, you MUST return 2 jobs.
    2. Use **markdown bolding** to highlight key skills and metrics within the bullet points.
    3. {tier_instructions}
    4. TRANSPARENCY: If you remove ANY bullet points, skills, or sentences from the Extracted JSON during your optimization, you MUST list them in `items_removed_for_optimization`.

    Return ONLY a valid JSON object matching this schema:
    {{
      "ats_match_score": 85,
      "missing_keywords": ["Skill 1", "Skill 2"],
      "items_removed_for_optimization": [
         {{"item": "Specific bullet removed", "reason": "Irrelevant to JD"}}
      ],
      "contact_info": {{"name": "...", "contact_string": "..."}},
      "summary": "A highly tailored professional summary. Use **bolding**.",
      "skills": ["**Skill:** Description..."],
      "experience": [
          {{
            "title": "...",
            "company": "...",
            "dates": "...",
            "bullets": ["**Action verb** describing achievement..."]
          }}
      ],
      "certifications": ["..."],
      "gap_bridger_project": "Project Title (if Tier 3, else empty string)",
      "gap_bridger_bullets": []
    }}
    """

    try:
        optimizer_response = coach_llm_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "system", "content": optimizer_prompt}],
            temperature=0.3,
            response_format={"type": "json_object"}
        )
        final_ai_data = json.loads(optimizer_response.choices[0].message.content)
        print(f"\n[STEP 2] Jobs Optimized: {len(final_ai_data.get('experience', []))}\n")

    except Exception as e:
        logger.error(f"Optimizer failed: {e}")
        # SECURE REFUND: Only refund if it's a paid tier and NOT a guest!
        if user_id != "guest" and tier > 1:
            await asyncio.to_thread(lambda: supabase.table("user_credits").update({"balance_minutes": curr_bal}).eq("user_id", user_id).execute())
        raise HTTPException(status_code=500, detail="AI Generation failed.")

    # 3. Generate the Word Document
    doc_io = create_optimized_word_doc(final_ai_data, final_resume_text)

    # 4. Return the file and headers
    headers = {
        'Content-Disposition': 'attachment; filename="Optimized_Resume.docx"',
        'X-ATS-Score': str(final_ai_data.get("ats_match_score", 0)),
        'X-Missing-Keywords': json.dumps(final_ai_data.get("missing_keywords", [])),
        'X-Items-Removed': json.dumps(final_ai_data.get("items_removed_for_optimization", [])),
        'Access-Control-Expose-Headers': 'X-ATS-Score, X-Missing-Keywords, X-Items-Removed' 
    }

    return StreamingResponse(doc_io, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers=headers)

# ==========================================
#         AI COACH ENDPOINTS
# ==========================================

@app.post("/coach/start")
async def start_coaching(
    user_id: str = Form(...),
    job_description: str = Form(""),
    difficulty: str = Form("Medium"),
    resume_text: str = Form(""),
    resume_file: UploadFile = File(None)
):
    print(f"--- STARTING COACH SESSION FOR USER: {user_id} ---")
    final_resume_text = resume_text

    if resume_file:
        print(f"Received file: {resume_file.filename}")
        try:
            content = await resume_file.read()
            final_resume_text = extract_text_from_file(content, resume_file.filename)
            print(f"Successfully extracted {len(final_resume_text)} characters from file.")
        except Exception as e:
            logger.error(f"File extraction failed: {str(e)}")
            return {"error": f"Could not read file: {str(e)}"}

    if not final_resume_text.strip():
        print("Warning: No resume text provided or extracted.")
        final_resume_text = "No resume provided. Ask general interview questions."

    diff_instruction = get_difficulty_instruction(difficulty)
    
    prompt = f"""
    You are an expert technical interviewer. 
    Difficulty Level: {difficulty}
    {diff_instruction}
    
    User Resume: {final_resume_text[:2000]}
    Job Description: {job_description[:2000]}
    
    Start the interview. Output JUST the opening question.
    """
    
    print("Calling Groq API...")
    try:
        response = coach_llm_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "system", "content": prompt}],
            temperature=0.7,
            max_tokens=200
        )
        print("Groq API returned successfully.")
        
        return {
            "message": response.choices[0].message.content,
            "extracted_resume": final_resume_text 
        }
    except Exception as e:
        logger.error(f"Groq API Error: {str(e)}")
        # Return a 500 error structure that the frontend can at least read
        from fastapi import HTTPException
        raise HTTPException(status_code=500, detail=f"AI Error: {str(e)}")

@app.post("/coach/reply")
async def reply_coaching(data: CoachReply):
    diff_instruction = get_difficulty_instruction(data.difficulty)

    # UPDATED: We explicitly demand a JSON structure.
    messages = [
        {"role": "system", "content": f"""
        You are a strict but helpful Interview Coach.
        Difficulty Level: {data.difficulty}
        {diff_instruction}
        
        Job Description: {data.job_description[:500]}
        Resume: {data.resume_text[:1000]}
        
        The user just answered your question.
        1. Analyze their answer against their resume.
        2. Give a short rating (Poor/Good/Excellent).
        3. Provide 1 sentence of specific feedback.
        4. Ask the NEXT question based on the difficulty level.
        
        CRITICAL: You MUST respond ONLY with a valid JSON object in the exact format below. Do not include markdown code blocks, just raw JSON.
        {{
            "rating": "[Poor/Good/Excellent]",
            "feedback": "Your 1 sentence feedback here.",
            "next_question": "Your next question here."
        }}
        """}
    ]
    
    for msg in data.history:
        messages.append(msg)
    
    messages.append({"role": "user", "content": data.user_answer})

    response = coach_llm_client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=messages,
        temperature=0.7,
        max_tokens=500,
        response_format={"type": "json_object"} # FORCING JSON OUTPUT (Groq supports this)
    )
    
    return {"message": response.choices[0].message.content}

@app.post("/coach/end")
async def end_coaching(data: CoachReply):
    messages = [
        {"role": "system", "content": f"""
        You are a strict but helpful Interview Coach. The interview is now OVER.
        Generate a final Scorecard based on the candidate's answers.
        Resume Context: {data.resume_text[:500]}
        
        CRITICAL: You MUST respond ONLY with a valid JSON object in the exact format below. 
        Do not ask any more questions. Do not include markdown code blocks.
        {{
            "overall_score": 75,
            "summary": "A 2-sentence overall summary of their performance highlighting their main strength.",
            "areas_of_improvement": [
                "First specific actionable bullet point.",
                "Second specific actionable bullet point.",
                "Third specific actionable bullet point."
            ]
        }}
        """}
    ]
    
    # Append the history so it knows what to grade
    for msg in data.history:
        messages.append(msg)

    response = coach_llm_client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=messages,
        temperature=0.7,
        max_tokens=600,
        response_format={"type": "json_object"} # Force strict JSON
    )
    return {"message": response.choices[0].message.content}


# ==========================================
#         WEBSOCKET (LIVE COPILOT)
# ==========================================

@app.websocket("/ws")
async def websocket_endpoint(websocket: WebSocket, token: str = None):
    await websocket.accept()
    logger.info("🚀 NEW CONNECTION ATTEMPT")

    if not token:
        logger.error("❌ No token provided. Closing.")
        await websocket.close(code=1008)
        return

    loop = asyncio.get_event_loop()

    try:
        user_res = await asyncio.to_thread(supabase.auth.get_user, token)
        user_id = user_res.user.id

        current_brain = get_brain_for_user(user_id)

        credit_res = await asyncio.to_thread(
            lambda: supabase.table("user_credits").select("balance_minutes").eq("user_id", user_id).single().execute()
        )
        balance = credit_res.data.get("balance_minutes", 0)

        if balance <= 0:
            logger.info(f"User {user_id} has 0 credits.")
            await websocket.send_json({"event": "out_of_credits"})
            await websocket.close()
            return

        logger.info(f"✅ User {user_id} Connected | Balance: {balance}m")

    except Exception as e:
        logger.error(f"❌ Auth Error: {e}")
        await websocket.close(code=1008)
        return

    countdown_active = True
    async def credit_countdown():
        while countdown_active:
            await asyncio.sleep(60)
            if not countdown_active: break
            
            try:
                curr_res = await asyncio.to_thread(
                    lambda: supabase.table("user_credits").select("balance_minutes").eq("user_id", user_id).single().execute()
                )
                curr_bal = curr_res.data.get("balance_minutes", 0)

                if curr_bal > 0:
                    new_bal = curr_bal - 1
                    await asyncio.to_thread(
                        lambda: supabase.table("user_credits").update({"balance_minutes": new_bal}).eq("user_id", user_id).execute()
                    )
                    await websocket.send_json({"event": "credit_update", "balance": new_bal})
                    
                    if new_bal <= 0:
                        await websocket.send_json({"event": "out_of_credits"})
                        await websocket.close()
                        break
            except Exception as e:
                logger.error(f"⚠️ Countdown error: {e}")

    countdown_task = asyncio.create_task(credit_countdown())

    try:
        config = DeepgramClientOptions(url="api.deepgram.com", verbose=logging.WARNING)
        deepgram = DeepgramClient(DEEPGRAM_API_KEY, config)
        dg_connection = deepgram.listen.live.v("1")
    except Exception as e:
        logger.error(f"Deepgram Init Failed: {e}")
        countdown_task.cancel()
        await websocket.close()
        return

    transcript_buffer = [] 
    
    async def trigger_ai_response(text):
        if len(text.strip()) < 2: return

        logger.info(f"🚀 AI Triggered for {user_id}: '{text[:30]}...'")
        await websocket.send_json({"event": "ai_start"})
        
        system_prompt = current_brain.build_system_prompt()
        messages = [{"role": "system", "content": system_prompt}]
        messages.extend(current_brain.history)
        messages.append({"role": "user", "content": text})

        full_answer = ""
        try:
            async for chunk in stream_completion(messages):
                full_answer += chunk
                await websocket.send_json({"event": "ai_chunk", "text": chunk})
            
            current_brain.add_interaction(text, full_answer)
            await websocket.send_json({"event": "ai_done"})
        except Exception as e:
            logger.error(f"AI Error: {e}")
            await websocket.send_json({"event": "ai_done"})

    # DEEPGRAM EVENT HANDLERS
    def on_message(self, result, **kwargs):
        sentence = result.channel.alternatives[0].transcript
        if len(sentence) > 0:
            is_final = result.is_final
            asyncio.run_coroutine_threadsafe(
                websocket.send_json({"event": "transcript", "text": sentence, "is_final": is_final}), loop
            )
            if is_final:
                transcript_buffer.append(sentence)
                if sentence.strip().endswith("?"):
                    full_text = " ".join(transcript_buffer)
                    transcript_buffer.clear()
                    asyncio.run_coroutine_threadsafe(trigger_ai_response(full_text), loop)

    def on_utterance_end(self, utterance_end, **kwargs):
        if len(transcript_buffer) > 0:
            full_text = " ".join(transcript_buffer)
            if len(full_text.split()) >= 2:
                transcript_buffer.clear()
                asyncio.run_coroutine_threadsafe(trigger_ai_response(full_text), loop)

    dg_connection.on(LiveTranscriptionEvents.Transcript, on_message)
    dg_connection.on(LiveTranscriptionEvents.UtteranceEnd, on_utterance_end)
    
    options = LiveOptions(
        model="nova-2", 
        language="en-US", 
        smart_format=True, 
        interim_results=True, 
        utterance_end_ms="1000", 
        vad_events=True
    )
    
    if dg_connection.start(options) is False:
        logger.error("Failed to connect to Deepgram")
        countdown_task.cancel()
        await websocket.close()
        return

    try:
        while True:
            message = await websocket.receive()
            if "bytes" in message:
                dg_connection.send(message["bytes"])
            if "text" in message:
                msg = json.loads(message["text"])
                if msg.get("text") == "stop": break
    except WebSocketDisconnect:
        logger.info(f"🔴 User {user_id} Disconnected")
    finally:
        countdown_active = False 
        countdown_task.cancel()
        dg_connection.finish()